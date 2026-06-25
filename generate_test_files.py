import os
import docx
from fpdf import FPDF

def create_docx():
    doc = docx.Document()
    
    # Title
    doc.add_heading('JOHN DOE', level=0)
    
    # Subtitle / Contact Info
    p = doc.add_paragraph()
    p.add_run('Email: john.doe@email.com | Phone: 555-019-2834 | Location: Austin, Texas\n')
    p.add_run('LinkedIn: linkedin.com/in/johndoepy | GitHub: github.com/johndoepy')
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph(
        "Reliable and motivated Software Developer with over 4 years of experience writing backend Python code. "
        "Familiar with web frameworks, database administration, and frontend technologies. Eager to join a "
        "growth-oriented company and contribute to engineering projects."
    )
    
    # Technical Skills
    doc.add_heading('TECHNICAL SKILLS', level=1)
    p_skills = doc.add_paragraph()
    p_skills.add_run('• Programming: ').bold = True
    p_skills.add_run('Python, SQL, JavaScript\n')
    p_skills.add_run('• Frameworks: ').bold = True
    p_skills.add_run('Flask, Django, Bootstrap, jQuery\n')
    p_skills.add_run('• Databases: ').bold = True
    p_skills.add_run('PostgreSQL, SQLite\n')
    p_skills.add_run('• Tools & Systems: ').bold = True
    p_skills.add_run('Git, Jira, Slack, Linux')
    
    # Work Experience
    doc.add_heading('WORK EXPERIENCE', level=1)
    
    doc.add_heading('Python Developer | TechCorp Solutions (Jan 2022 - Present)', level=2)
    doc.add_paragraph('• Worked on backend Python code for a customer management web portal.', style='List Bullet')
    doc.add_paragraph('• Responsible for writing and maintaining database queries and schema migrations in PostgreSQL.', style='List Bullet')
    doc.add_paragraph('• Assisted with troubleshooting issues and fixing bugs reported by clients.', style='List Bullet')
    doc.add_paragraph('• Collaborated with frontend team members to connect endpoints to Bootstrap templates.', style='List Bullet')
    doc.add_paragraph('• Attended weekly Agile Scrum standup meetings.', style='List Bullet')
    
    doc.add_heading('Software Associate | WebDev Builders (Sep 2020 - Dec 2021)', level=2)
    doc.add_paragraph('• Developed Flask web applications for local business clients.', style='List Bullet')
    doc.add_paragraph('• Wrote unit tests for key application features to check for errors.', style='List Bullet')
    doc.add_paragraph('• Handled deployment of websites to Heroku.', style='List Bullet')
    doc.add_paragraph('• Managed user authentication and form validation code.', style='List Bullet')
    
    # Education
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science (2016 - 2020)\nUniversity of Texas, Austin')
    
    # Projects
    doc.add_heading('PROJECTS', level=1)
    
    doc.add_heading('Task Tracker App', level=2)
    doc.add_paragraph("Developed a simple web application using Flask and SQLite to help users organize daily tasks. Added user login pages and database tracking.")
    
    doc.add_heading('Weather Dashboard', level=2)
    doc.add_paragraph("Created a frontend page that calls a weather API and displays local forecasts using jQuery.")
    
    # Save docx
    docx_path = "sample_resume.docx"
    doc.save(docx_path)
    print(f"Created: {docx_path}")

def create_pdf():
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    
    # Fonts
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 8, "JOHN DOE", align="C", new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 5, "Email: john.doe@email.com | Phone: 555-019-2834 | Location: Austin, Texas", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, "LinkedIn: linkedin.com/in/johndoepy | GitHub: github.com/johndoepy", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    
    # Sections
    sections = [
        ("PROFESSIONAL SUMMARY", [
            "Reliable and motivated Software Developer with over 4 years of experience writing backend Python code. "
            "Familiar with web frameworks, database administration, and frontend technologies. Eager to join a "
            "growth-oriented company and contribute to engineering projects."
        ]),
        ("TECHNICAL SKILLS", [
            "Programming: Python, SQL, JavaScript",
            "Frameworks: Flask, Django, Bootstrap, jQuery",
            "Databases: PostgreSQL, SQLite",
            "Tools & Systems: Git, Jira, Slack, Linux"
        ]),
        ("WORK EXPERIENCE", [
            "Python Developer | TechCorp Solutions (Jan 2022 - Present)",
            "  * Worked on backend Python code for a customer management web portal.",
            "  * Responsible for writing and maintaining database queries and schema migrations in PostgreSQL.",
            "  * Assisted with troubleshooting issues and fixing bugs reported by clients.",
            "  * Collaborated with frontend team members to connect endpoints to Bootstrap templates.",
            "  * Attended weekly Agile Scrum standup meetings.",
            "",
            "Software Associate | WebDev Builders (Sep 2020 - Dec 2021)",
            "  * Developed Flask web applications for local business clients.",
            "  * Wrote unit tests for key application features to check for errors.",
            "  * Handled deployment of websites to Heroku.",
            "  * Managed user authentication and form validation code."
        ]),
        ("EDUCATION", [
            "Bachelor of Science in Computer Science (2016 - 2020)",
            "University of Texas, Austin"
        ]),
        ("PROJECTS", [
            "Task Tracker App",
            "  * Developed a simple web application using Flask and SQLite to help users organize daily tasks.",
            "  * Added user login pages and database tracking.",
            "",
            "Weather Dashboard",
            "  * Created a frontend page that calls a weather API and displays local forecasts using jQuery."
        ])
    ]
    
    for title, lines in sections:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 6, title, align="L", new_x="LMARGIN", new_y="NEXT")
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
        pdf.ln(2)
        
        pdf.set_font("Helvetica", "", 10)
        for line in lines:
            if not line.strip():
                pdf.ln(3)
                continue
            print(f"DEBUG Line: '{line[:20]}...' at x={pdf.get_x():.2f}, y={pdf.get_y():.2f}")
            if line.startswith("  *") or line.startswith("•"):
                pdf.multi_cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
            elif " | " in line and ("Present" in line or "2021" in line):
                pdf.set_font("Helvetica", "B", 10)
                pdf.cell(0, 5, line, align="L", new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Helvetica", "", 10)
            else:
                pdf.multi_cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
        
    pdf_path = "sample_resume.pdf"
    pdf.output(pdf_path)
    print(f"Created: {pdf_path}")

if __name__ == "__main__":
    create_docx()
    create_pdf()
